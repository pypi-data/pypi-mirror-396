import os
import re
from datetime import datetime
from markdown import markdown
from bs4 import BeautifulSoup
from weasyprint import HTML
import openpyxl
from openpyxl.styles import Font, Border, Side, Alignment

from src import terminal as term


def has_table(content: str) -> bool:
    return bool(re.search(r'\|.+\|.*\n\|[-:| ]+\|', content))


def get_filename(ext: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    exports_dir = os.path.join(os.getcwd(), "exports")
    os.makedirs(exports_dir, exist_ok=True)
    return os.path.join(exports_dir, f"tenk_{timestamp}.{ext}")


def md_to_html(content: str) -> str:
    html = markdown(content, extensions=['tables', 'fenced_code'])
    return f"""
    <html>
    <head>
        <style>
            @page {{ margin: 0.75in; }}
            body {{ font-family: 'Times New Roman', Times, serif; font-size: 12pt; line-height: 1.4; margin: 0; }}
            h1 {{ font-size: 18pt; margin: 0 0 8pt 0; }}
            h2 {{ font-size: 14pt; margin: 12pt 0 6pt 0; }}
            h3 {{ font-size: 12pt; margin: 10pt 0 4pt 0; }}
            p {{ margin: 0 0 8pt 0; }}
            code {{ font-family: 'Courier New', monospace; font-size: 10pt; background: #f5f5f5; padding: 1px 3px; }}
            pre {{ background: #f5f5f5; padding: 8pt; margin: 8pt 0; font-size: 10pt; }}
            pre code {{ background: none; padding: 0; }}
            table {{ border-collapse: collapse; width: 100%; margin: 8pt 0; }}
            th, td {{ border: 1px solid #000; padding: 4pt 6pt; text-align: left; font-size: 11pt; }}
            th {{ font-weight: bold; }}
            blockquote {{ margin: 8pt 0 8pt 20pt; padding-left: 8pt; border-left: 2px solid #666; }}
            ul, ol {{ margin: 4pt 0; padding-left: 20pt; }}
            li {{ margin: 2pt 0; }}
        </style>
    </head>
    <body>{html}</body>
    </html>
    """


def export_pdf(content: str) -> str:
    filepath = get_filename("pdf")
    html = md_to_html(content)
    HTML(string=html).write_pdf(filepath)
    return filepath


def export_docx(content: str) -> str:
    from docx import Document
    from docx.shared import Pt, Inches

    filepath = get_filename("docx")
    doc = Document()

    html = markdown(content, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')

    for elem in soup.children:
        if elem.name == 'h1':
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == 'h2':
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == 'h3':
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == 'p':
            p = doc.add_paragraph()
            add_inline(p, elem)
        elif elem.name == 'ul':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                add_inline(p, li)
        elif elem.name == 'ol':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                add_inline(p, li)
        elif elem.name == 'pre':
            code = elem.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif elem.name == 'blockquote':
            p = doc.add_paragraph(elem.get_text())
            p.paragraph_format.left_indent = Inches(0.5)
        elif elem.name == 'table':
            rows = elem.find_all('tr')
            if rows:
                cols = len(rows[0].find_all(['th', 'td']))
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        table.rows[i].cells[j].text = cell.get_text()

    doc.save(filepath)
    return filepath


def add_inline(paragraph, elem):
    for child in elem.children if hasattr(elem, 'children') else [elem]:
        if isinstance(child, str):
            paragraph.add_run(child)
        elif child.name == 'strong' or child.name == 'b':
            paragraph.add_run(child.get_text()).bold = True
        elif child.name == 'em' or child.name == 'i':
            paragraph.add_run(child.get_text()).italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
        elif child.name == 'a':
            paragraph.add_run(child.get_text())
        else:
            paragraph.add_run(child.get_text() if hasattr(child, 'get_text') else str(child))


def export_xlsx(content: str) -> str:
    filepath = get_filename("xlsx")
    wb = openpyxl.Workbook()

    html = markdown(content, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')
    tables = soup.find_all('table')

    if not tables:
        return None

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    for idx, table in enumerate(tables):
        if idx == 0:
            ws = wb.active
            ws.title = f"Table {idx + 1}"
        else:
            ws = wb.create_sheet(title=f"Table {idx + 1}")

        rows = table.find_all('tr')
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for col_idx, cell in enumerate(cells):
                excel_cell = ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell.get_text().strip())
                excel_cell.border = thin_border
                excel_cell.alignment = Alignment(wrap_text=True, vertical='top')
                if cell.name == 'th':
                    excel_cell.font = Font(bold=True)

        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max_len + 2, 50)

    wb.save(filepath)
    return filepath


def handle_export(choice: str, content: str) -> str:
    exporters = {
        'pdf': export_pdf,
        'docx': export_docx,
        'xlsx': export_xlsx,
    }

    if choice.lower() not in exporters:
        return None

    try:
        filepath = exporters[choice.lower()](content)
        return filepath
    except Exception as e:
        term.print_error(f"Export failed: {e}")
        return None
