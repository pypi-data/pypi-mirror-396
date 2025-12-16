"""
Word文档内容读取工具的辅助函数

提供文档处理的基础工具函数
"""

import os
import zipfile
import shutil
import re
from typing import Dict, List, Any
from docx import Document


def ensure_docx_extension(filename: str) -> str:
    """确保文件名有.docx扩展名"""
    if not filename.endswith('.docx'):
        filename += '.docx'
    return filename


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """获取Word文档的属性信息"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document {doc_path} does not exist")
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        # 计算更准确的字数统计（支持中文）
        total_chars = 0
        total_words = 0
        all_text = ""
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # 跳过空段落
                all_text += text + " "
                total_chars += len(text.replace(' ', ''))  # 不包括空格的字符数
                # 对于包含中文的文本，按字符计算；对于英文，按单词计算
                if any('\u4e00' <= char <= '\u9fff' for char in text):
                    # 包含中文字符，按字符数计算
                    total_words += len([char for char in text if char.strip()])
                else:
                    # 纯英文，按单词计算
                    total_words += len(text.split())
        
        # 同时处理表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            all_text += text + " "
                            total_chars += len(text.replace(' ', ''))
                            if any('\u4e00' <= char <= '\u9fff' for char in text):
                                total_words += len([char for char in text if char.strip()])
                            else:
                                total_words += len(text.split())
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "word_count": total_words,
            "character_count": total_chars,  # 添加字符数统计
            "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),  # 只计算非空段落
            "section_count": len(doc.sections),  # 添加节数统计
            "table_count": len(doc.tables),
            "document_language": "mixed" if any('\u4e00' <= char <= '\u9fff' for char in all_text) else "english"  # 简单的语言检测
        }
    except Exception as e:
        raise RuntimeError(f"Failed to get document properties: {str(e)}")


def extract_document_text(doc_path: str) -> str:
    """从Word文档中提取所有文本"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document {doc_path} does not exist")
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text: {str(e)}")


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """获取Word文档的结构信息"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document {doc_path} does not exist")
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # 获取段落信息
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # 获取表格信息
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # 获取表格数据样本
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        raise RuntimeError(f"Failed to get document structure: {str(e)}")


def get_paragraph_text(doc_path: str, paragraph_index: int) -> Dict[str, Any]:
    """获取指定段落的文本内容"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document {doc_path} does not exist")
    
    try:
        doc = Document(doc_path)
        
        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            raise ValueError(f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs.")
        
        paragraph = doc.paragraphs[paragraph_index]
        
        return {
            "index": paragraph_index,
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "is_heading": paragraph.style.name.startswith("Heading") if paragraph.style else False
        }
    except Exception as e:
        raise RuntimeError(f"Failed to get paragraph text: {str(e)}")


def find_text(doc_path: str, text_to_find: str, match_case: bool = True, whole_word: bool = False) -> Dict[str, Any]:
    """在Word文档中查找指定文本的所有出现位置"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document {doc_path} does not exist")
    
    if not text_to_find:
        raise ValueError("Search text cannot be empty")
    
    try:
        doc = Document(doc_path)
        results = {
            "query": text_to_find,
            "match_case": match_case,
            "whole_word": whole_word,
            "occurrences": [],
            "total_count": 0
        }
        
        search_text = text_to_find if match_case else text_to_find.lower()
        
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text if match_case else para.text.lower()
            
            # 查找所有出现位置
            start_pos = 0
            while True:
                if whole_word:
                    # 全词匹配搜索
                    words = para_text.split()
                    found = False
                    for word_idx, word in enumerate(words):
                        if (word == search_text or 
                            (not match_case and word.lower() == search_text.lower())):
                            results["occurrences"].append({
                                "paragraph_index": i,
                                "position": word_idx,
                                "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                            })
                            results["total_count"] += 1
                            found = True
                    
                    # 检查完所有单词后跳出
                    break
                else:
                    # 子字符串搜索
                    pos = para_text.find(search_text, start_pos)
                    if pos == -1:
                        break
                    
                    results["occurrences"].append({
                        "paragraph_index": i,
                        "position": pos,
                        "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                    })
                    results["total_count"] += 1
                    start_pos = pos + 1
        
        return results
    except Exception as e:
        raise RuntimeError(f"Failed to search text: {str(e)}")


def extract_images_from_docx(doc_path: str, output_dir: str) -> Dict[str, Any]:
    """从Word文档中提取所有图片到指定目录"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document {doc_path} does not exist")
        
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    extracted_count = 0
    extracted_files = []
    
    # 定义允许的图片扩展名
    VALID_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff', '.svg', '.wmf', '.emf', '.ico', '.webp'}
    
    try:
        with zipfile.ZipFile(doc_path) as doc_zip:
            # 获取所有媒体文件
            all_media_files = [f for f in doc_zip.namelist() if f.startswith('word/media/')]
            
            # 过滤有效图片文件
            media_files = []
            for f in all_media_files:
                ext = os.path.splitext(f)[1].lower()
                if ext in VALID_EXTENSIONS:
                    media_files.append(f)
            
            # 尝试按照文件名中的数字排序
            def get_file_number(filename):
                match = re.search(r'image(\d+)\.', filename)
                return int(match.group(1)) if match else 999999
            
            media_files.sort(key=get_file_number)
            
            for i, media_file in enumerate(media_files):
                # 获取扩展名
                ext = os.path.splitext(media_file)[1] # 保持原始大小写
                
                # 构建新文件名，保持顺序
                new_filename = f"image_{i+1}{ext}"
                target_path = os.path.join(output_dir, new_filename)
                
                # 读取并写入
                with doc_zip.open(media_file) as source, open(target_path, 'wb') as target:
                    shutil.copyfileobj(source, target)
                    
                extracted_files.append(new_filename)
                extracted_count += 1
                
        return {
            "total_images": extracted_count,
            "output_dir": output_dir,
            "images": extracted_files
        }
    except Exception as e:
        raise RuntimeError(f"Failed to extract images: {str(e)}")
