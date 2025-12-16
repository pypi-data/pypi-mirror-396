"""Microsoft Word document loader."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from prompt_amplifier.core.exceptions import LoaderError
from prompt_amplifier.loaders.base import BaseLoader
from prompt_amplifier.models.document import Document


class DocxLoader(BaseLoader):
    """
    Load Microsoft Word documents (.docx).

    Requires: python-docx

    Example:
        >>> loader = DocxLoader()
        >>> docs = loader.load("document.docx")
    """

    def __init__(self, **kwargs: Any):
        super().__init__(**kwargs)
        self._check_dependency()

    def _check_dependency(self) -> None:
        """Check if python-docx is installed."""
        try:
            import docx  # noqa: F401
        except ImportError:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                "Install it with: pip install python-docx"
            )

    def load(self, source: str | Path) -> list[Document]:
        """Load a Word document."""
        import docx

        path = Path(source)

        if not path.exists():
            raise LoaderError(f"File not found: {source}")

        try:
            doc = docx.Document(str(path))
        except Exception as e:
            raise LoaderError(f"Failed to read Word document {source}: {e}")

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_text))
            if table_rows:
                tables_text.append("\n".join(table_rows))

        # Combine all content
        content_parts = paragraphs
        if tables_text:
            content_parts.append("\n\n--- Tables ---\n")
            content_parts.extend(tables_text)

        content = "\n\n".join(content_parts)

        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
        }

        return [
            self._create_document(
                content=content,
                source=source,
                source_type="docx",
                metadata=metadata,
            )
        ]

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]
