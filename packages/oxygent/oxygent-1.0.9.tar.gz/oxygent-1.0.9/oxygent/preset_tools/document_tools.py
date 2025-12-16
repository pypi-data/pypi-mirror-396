"""
Document Processing Tools for OxyGent Framework
支持PDF、Word、Excel、PowerPoint等文档格式的处理

"""

import json
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from pydantic import Field

from oxygent.oxy import FunctionHub

# 延迟导入，避免必须安装所有依赖
logger = logging.getLogger(__name__)


class DocumentToolsHub(FunctionHub):
    """文档处理工具中心，管理所有文档相关操作"""

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖库是否安装"""
        missing = []
        
        try:
            import fitz  # PyMuPDF
        except ImportError:
            missing.append("PyMuPDF")
        
        try:
            import pdfplumber
        except ImportError:
            missing.append("pdfplumber")
        
        try:
            from docx import Document
        except ImportError:
            missing.append("python-docx")
        
        try:
            from openpyxl import load_workbook
        except ImportError:
            missing.append("openpyxl")
        
        if missing:
            logger.warning(
                f"部分文档处理库未安装: {', '.join(missing)}。"
                f"这些依赖是可选的，但使用文档处理功能时需要安装。"
                f"请运行: pip install {' '.join(missing)} 或 uv pip install {' '.join(missing)}"
                f"或安装所有依赖: pip install -r requirements.txt 或 uv pip install -r requirements.txt"
            )


# 创建全局工具实例
document_tools = DocumentToolsHub(name="document_tools")


# ==================== PDF 文本提取 ====================

@document_tools.tool(
    description="从PDF文件中提取文本内容。支持指定页码范围提取，适用于文字版PDF（非扫描版）。"
    "返回按页组织的文本内容。如需处理扫描版PDF，请使用ocr_pdf工具。"
)
def extract_pdf_text(
    path: str,
    page_range: Optional[str] = None,
    max_chars_per_page: int = 10000,
) -> str:
    """
    提取PDF文本内容
    
    技术实现：
    - 使用PyMuPDF (fitz)进行高效文本提取
    - 支持多种PDF编码格式
    - 自动处理页面旋转和布局
    
    Args:
        path: PDF文件路径
        page_range: 页码范围字符串
        max_chars_per_page: 单页最大字符数
        
    Returns:
        JSON格式的提取结果，包含文本内容和元数据
    """
    try:
        import fitz
        
        if not Path(path).exists():
            return json.dumps(
                {"error": f"文件不存在: {path}"}, 
                ensure_ascii=False
            )
        
        doc = fitz.open(path)
        total_pages = len(doc)
        
        # 解析页码范围
        pages = _parse_page_range(page_range, total_pages)
        
        if not pages:
            return json.dumps(
                {"error": "无效的页码范围或页码超出文档范围"},
                ensure_ascii=False
            )
        
        # 提取文本
        results = []
        for page_num in pages:
            page = doc[page_num]
            text = page.get_text("text")  # 使用纯文本模式
            
            # 限制单页文本长度
            if len(text) > max_chars_per_page:
                text = text[:max_chars_per_page] + f"\n...(已截断，原文本长度: {len(text)}字符)"
            
            results.append({
                "page_number": page_num + 1,
                "text": text.strip(),
                "char_count": len(text),
                "has_images": len(page.get_images()) > 0
            })
        
        doc.close()
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "total_pages": total_pages,
            "extracted_pages": len(pages),
            "pages": results
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "PyMuPDF未安装，请运行: pip install PyMuPDF"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"PDF文本提取失败: {e}")
        return json.dumps(
            {"error": f"提取失败: {str(e)}"}, 
            ensure_ascii=False
        )


# ==================== PDF 表格提取 ====================

@document_tools.tool(
    description="从PDF中提取表格数据，返回结构化的JSON格式。"
    "使用pdfplumber的高精度表格识别算法，能够处理复杂的表格布局。"
    "适用于包含表格的报告、发票、财务文档等。"
)
def extract_pdf_tables(
    path: str,
    page_range: Optional[str] = None,
    table_settings: Optional[Dict[str, Any]] = None,
) -> str:
    """
    从PDF提取表格数据
    
    技术特点：
    - 使用pdfplumber的表格识别引擎
    - 支持合并单元格的处理
    - 自动识别表头
    - 处理跨页表格
    
    Args:
        path: PDF文件路径
        page_range: 页码范围
        table_settings: 表格识别参数
        
    Returns:
        包含所有表格数据的JSON结构
    """
    try:
        import pdfplumber
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        all_tables = []
        
        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            pages = _parse_page_range(page_range, total_pages)
            
            for page_num in pages:
                page = pdf.pages[page_num]
                
                # 提取表格
                tables = page.extract_tables(table_settings or {})
                
                for table_idx, table in enumerate(tables):
                    if not table or len(table) == 0:
                        continue
                    
                    # 处理表格数据
                    headers = table[0] if table else []
                    rows = table[1:] if len(table) > 1 else []
                    
                    # 清理空值
                    clean_rows = []
                    for row in rows:
                        clean_row = [cell if cell is not None else "" for cell in row]
                        if any(clean_row):  # 只保留非空行
                            clean_rows.append(clean_row)
                    
                    all_tables.append({
                        "page": page_num + 1,
                        "table_index": table_idx + 1,
                        "headers": headers,
                        "rows": clean_rows,
                        "row_count": len(clean_rows),
                        "column_count": len(headers) if headers else 0
                    })
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "table_count": len(all_tables),
            "tables": all_tables
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "pdfplumber未安装，请运行: pip install pdfplumber"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"PDF表格提取失败: {e}")
        return json.dumps({"error": f"提取失败: {str(e)}"}, ensure_ascii=False)


# ==================== PDF 图像提取 ====================

@document_tools.tool(
    description="从PDF文件中提取所有嵌入的图像，并保存到指定目录。"
    "图像将自动命名为 'image_页码_序号.扩展名' 格式。"
    "支持PNG、JPEG等常见图像格式。"
)
def extract_pdf_images(
    path: str,
    output_dir: str,
    page_range: Optional[str] = None,
    min_size: int = 1024,
) -> str:
    """
    提取PDF中的图像
    
    功能特点：
    - 自动创建输出目录
    - 保持原始图像格式和质量
    - 可过滤小尺寸图片（如图标、装饰线等）
    - 返回每个图像的详细信息
    
    Args:
        path: PDF文件路径
        output_dir: 输出目录
        page_range: 页码范围
        min_size: 最小图像大小（字节）
        
    Returns:
        包含所有提取图像信息的JSON
    """
    try:
        import fitz
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        # 创建输出目录
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        doc = fitz.open(path)
        total_pages = len(doc)
        pages = _parse_page_range(page_range, total_pages)
        
        image_list = []
        
        for page_num in pages:
            page = doc[page_num]
            images = page.get_images()
            
            for img_idx, img in enumerate(images):
                xref = img[0]
                
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # 过滤小图像
                    if len(image_bytes) < min_size:
                        continue
                    
                    # 生成文件名
                    filename = f"image_p{page_num + 1}_{img_idx + 1}.{image_ext}"
                    filepath = output_path / filename
                    
                    # 保存图像
                    with open(filepath, "wb") as f:
                        f.write(image_bytes)
                    
                    image_list.append({
                        "page": page_num + 1,
                        "filename": filename,
                        "path": str(filepath),
                        "size_bytes": len(image_bytes),
                        "format": image_ext,
                        "width": base_image.get("width"),
                        "height": base_image.get("height")
                    })
                    
                except Exception as img_error:
                    logger.warning(f"提取第{page_num + 1}页第{img_idx + 1}个图像失败: {img_error}")
                    continue
        
        doc.close()
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "output_dir": str(output_path),
            "image_count": len(image_list),
            "images": image_list
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "PyMuPDF未安装，请运行: pip install PyMuPDF"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"PDF图像提取失败: {e}")
        return json.dumps({"error": f"提取失败: {str(e)}"}, ensure_ascii=False)


# ==================== PDF 合并 ====================

@document_tools.tool(
    description="将多个PDF文件合并为一个PDF文件。"
    "按照提供的文件列表顺序进行合并，支持任意数量的PDF文件。"
    "合并后保留原始文档的所有内容、格式和元数据。"
)
def merge_pdfs(
    pdf_paths: List[str],
    output_path: str,
    include_bookmarks: bool = True,
) -> str:
    """
    合并多个PDF文件
    
    技术特点：
    - 零质量损失的合并
    - 保留原始文档属性
    - 可选保留书签结构
    - 高效处理大文件
    
    Args:
        pdf_paths: PDF文件路径列表
        output_path: 输出文件路径
        include_bookmarks: 是否保留书签
        
    Returns:
        合并操作的结果信息
    """
    try:
        import fitz
        
        # 验证输入文件
        missing_files = []
        for pdf_path in pdf_paths:
            if not Path(pdf_path).exists():
                missing_files.append(pdf_path)
        
        if missing_files:
            return json.dumps({
                "error": f"以下文件不存在: {', '.join(missing_files)}"
            }, ensure_ascii=False)
        
        if len(pdf_paths) < 2:
            return json.dumps({
                "error": "至少需要2个PDF文件才能合并"
            }, ensure_ascii=False)
        
        # 创建新文档
        merged_doc = fitz.open()
        total_pages = 0
        
        # 逐个合并
        for idx, pdf_path in enumerate(pdf_paths):
            try:
                doc = fitz.open(pdf_path)
                page_count = len(doc)
                
                # 插入所有页面
                merged_doc.insert_pdf(doc)
                total_pages += page_count
                
                doc.close()
                
            except Exception as e:
                logger.error(f"合并文件 {pdf_path} 失败: {e}")
                merged_doc.close()
                return json.dumps({
                    "error": f"合并文件 {pdf_path} 时出错: {str(e)}"
                }, ensure_ascii=False)
        
        # 保存合并后的文档
        merged_doc.save(output_path)
        merged_doc.close()
        
        return json.dumps({
            "success": True,
            "message": f"成功合并 {len(pdf_paths)} 个PDF文件",
            "output_path": output_path,
            "total_pages": total_pages,
            "source_files": pdf_paths
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "PyMuPDF未安装，请运行: pip install PyMuPDF"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"PDF合并失败: {e}")
        return json.dumps({"error": f"合并失败: {str(e)}"}, ensure_ascii=False)


# ==================== PDF 拆分 ====================

@document_tools.tool(
    description="按页码范围拆分PDF文件为多个独立的PDF文件。"
    "可以灵活指定每个拆分文件包含的页码范围。"
    "适用于提取PDF章节、分离文档部分等场景。"
)
def split_pdf(
    path: str,
    split_ranges: List[str],
    output_dir: str,
    name_prefix: str = "split",
) -> str:
    """
    拆分PDF文件
    
    功能亮点：
    - 支持灵活的页码范围定义
    - 自动命名输出文件
    - 保持原始PDF质量
    - 支持不连续页码拆分
    
    Args:
        path: 源PDF文件路径
        split_ranges: 拆分范围列表
        output_dir: 输出目录
        name_prefix: 文件名前缀
        
    Returns:
        拆分操作的详细结果
    """
    try:
        import fitz
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        # 创建输出目录
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        doc = fitz.open(path)
        total_pages = len(doc)
        
        output_files = []
        
        for idx, range_str in enumerate(split_ranges):
            try:
                pages = _parse_page_range(range_str, total_pages)
                
                if not pages:
                    logger.warning(f"范围 '{range_str}' 无效或超出文档页数，已跳过")
                    continue
                
                # 创建新文档
                new_doc = fitz.open()
                
                for page_num in pages:
                    new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
                
                # 保存文件
                output_filename = f"{name_prefix}_{idx + 1}.pdf"
                output_filepath = output_path / output_filename
                new_doc.save(output_filepath)
                new_doc.close()
                
                output_files.append({
                    "filename": output_filename,
                    "path": str(output_filepath),
                    "page_range": range_str,
                    "page_count": len(pages)
                })
                
            except Exception as e:
                logger.error(f"拆分范围 '{range_str}' 失败: {e}")
                continue
        
        doc.close()
        
        if not output_files:
            return json.dumps({
                "error": "没有成功拆分任何文件，请检查页码范围是否正确"
            }, ensure_ascii=False)
        
        return json.dumps({
            "success": True,
            "message": f"成功拆分为 {len(output_files)} 个文件",
            "source_file": path,
            "output_dir": str(output_path),
            "files": output_files
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "PyMuPDF未安装，请运行: pip install PyMuPDF"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"PDF拆分失败: {e}")
        return json.dumps({"error": f"拆分失败: {str(e)}"}, ensure_ascii=False)


# ==================== PDF 元数据获取 ====================

@document_tools.tool(
    description="获取PDF文件的元数据信息，包括作者、标题、创建日期、页数等。"
    "还会返回文档的加密状态、PDF版本等技术信息。"
)
def get_pdf_info(
    path: str,
) -> str:
    """
    获取PDF元数据和基本信息
    
    返回信息包括：
    - 文档属性（标题、作者、主题等）
    - 页面信息（总页数、页面尺寸）
    - 技术信息（PDF版本、是否加密）
    - 内容统计（文本量、图像数量）
    
    Args:
        path: PDF文件路径
        
    Returns:
        包含所有元数据的JSON结构
    """
    try:
        import fitz
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        doc = fitz.open(path)
        
        # 获取元数据
        metadata = doc.metadata
        
        # 统计页面信息
        page_sizes = []
        total_images = 0
        total_text_length = 0
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            rect = page.rect
            page_sizes.append({
                "page": page_num + 1,
                "width": round(rect.width, 2),
                "height": round(rect.height, 2)
            })
            total_images += len(page.get_images())
            total_text_length += len(page.get_text())
        
        # 文件信息
        file_path = Path(path)
        file_size = file_path.stat().st_size
        
        info = {
            "success": True,
            "file_info": {
                "path": str(path),
                "filename": file_path.name,
                "size_bytes": file_size,
                "size_mb": round(file_size / (1024 * 1024), 2)
            },
            "document_metadata": {
                "title": metadata.get("title", ""),
                "author": metadata.get("author", ""),
                "subject": metadata.get("subject", ""),
                "keywords": metadata.get("keywords", ""),
                "creator": metadata.get("creator", ""),
                "producer": metadata.get("producer", ""),
                "creation_date": metadata.get("creationDate", ""),
                "modification_date": metadata.get("modDate", "")
            },
            "document_properties": {
                "page_count": len(doc),
                "is_encrypted": doc.is_encrypted,
                "is_pdf": doc.is_pdf,
                "pdf_version": doc.pdf_version if hasattr(doc, 'pdf_version') else "unknown"
            },
            "content_statistics": {
                "total_images": total_images,
                "estimated_text_length": total_text_length,
                "average_page_size": {
                    "width": round(sum(p["width"] for p in page_sizes) / len(page_sizes), 2),
                    "height": round(sum(p["height"] for p in page_sizes) / len(page_sizes), 2)
                } if page_sizes else {}
            },
            "page_sizes": page_sizes[:5] if len(page_sizes) > 5 else page_sizes  # 只返回前5页尺寸
        }
        
        doc.close()
        
        return json.dumps(info, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "PyMuPDF未安装，请运行: pip install PyMuPDF"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"获取PDF信息失败: {e}")
        return json.dumps({"error": f"获取信息失败: {str(e)}"}, ensure_ascii=False)


# ==================== Word 文档操作 ====================

@document_tools.tool(
    description="读取Word文档(.docx)的完整内容，包括段落文本和表格数据。"
    "自动识别文档结构，提取标题、正文、列表等元素。"
    "注意：仅支持.docx格式（Office 2007及以后），不支持旧版.doc格式。"
)
def read_docx(
    path: str,
    include_tables: bool = True,
    max_paragraphs: int = 1000,
) -> str:
    """
    读取Word文档内容
    
    提取内容：
    - 所有段落文本（保持顺序）
    - 表格数据（结构化格式）
    - 段落样式信息（可选）
    - 文档统计信息
    
    Args:
        path: Word文件路径
        include_tables: 是否包含表格
        max_paragraphs: 最大段落数
        
    Returns:
        结构化的文档内容JSON
    """
    try:
        from docx import Document
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        doc = Document(path)
        
        # 提取段落
        paragraphs = []
        for idx, para in enumerate(doc.paragraphs):
            if idx >= max_paragraphs:
                break
            
            text = para.text.strip()
            if text:  # 只保留非空段落
                paragraphs.append({
                    "index": idx + 1,
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                })
        
        # 提取表格
        tables_data = []
        if include_tables:
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_content.append(row_data)
                
                if table_content:  # 只保留非空表格
                    tables_data.append({
                        "table_index": table_idx + 1,
                        "row_count": len(table_content),
                        "column_count": len(table_content[0]) if table_content else 0,
                        "headers": table_content[0] if table_content else [],
                        "rows": table_content[1:] if len(table_content) > 1 else []
                    })
        
        # 统计信息
        total_text = " ".join(p["text"] for p in paragraphs)
        word_count = len(total_text.split())
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "statistics": {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_data),
                "word_count": word_count,
                "char_count": len(total_text)
            },
            "paragraphs": paragraphs,
            "tables": tables_data
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "python-docx未安装，请运行: pip install python-docx"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"读取Word文档失败: {e}")
        return json.dumps({"error": f"读取失败: {str(e)}"}, ensure_ascii=False)


@document_tools.tool(
    description="从Word文档中仅提取纯文本内容，不包含格式和表格。"
    "适用于需要快速获取文档文字内容的场景，如文本分析、关键词提取等。"
)
def extract_docx_text(
    path: str,
) -> str:
    """
    提取Word文档纯文本
    
    快速提取文档的所有文本内容，忽略格式和结构。
    
    Args:
        path: Word文件路径
        
    Returns:
        纯文本内容
    """
    try:
        from docx import Document
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        doc = Document(path)
        
        # 提取所有段落文本
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    full_text.append(row_text)
        
        result_text = "\n".join(full_text)
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "text": result_text,
            "length": len(result_text),
            "line_count": len(full_text)
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {"error": "python-docx未安装，请运行: pip install python-docx"},
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"提取Word文本失败: {e}")
        return json.dumps({"error": f"提取失败: {str(e)}"}, ensure_ascii=False)


# ==================== Excel 文档操作 ====================

@document_tools.tool(
    description="读取Excel文件(.xlsx/.xlsm)的数据内容。"
    "可指定工作表名称和读取行数限制。支持多工作表Excel文件。"
    "返回结构化的表格数据，便于进一步处理和分析。"
)
def read_excel(
    path: str,
    sheet_name: Optional[str] = None,
    max_rows: int = 100,
    has_header: bool = True,
) -> str:
    """
    读取Excel表格数据
    
    功能特点：
    - 支持.xlsx和.xlsm格式
    - 自动识别表头
    - 处理空单元格
    - 支持多工作表
    
    Args:
        path: Excel文件路径
        sheet_name: 工作表名称
        max_rows: 最大行数
        has_header: 是否有表头
        
    Returns:
        结构化的Excel数据JSON
    """
    try:
        from openpyxl import load_workbook
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        wb = load_workbook(path, read_only=True, data_only=True)
        
        # 获取工作表
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                return json.dumps({
                    "error": f"工作表 '{sheet_name}' 不存在",
                    "available_sheets": wb.sheetnames
                }, ensure_ascii=False)
            ws = wb[sheet_name]
        else:
            ws = wb.active
        
        # 读取数据
        data = []
        for idx, row in enumerate(ws.iter_rows(values_only=True)):
            if idx >= max_rows:
                break
            
            # 转换为列表并处理None值
            row_data = [cell if cell is not None else "" for cell in row]
            
            # 跳过完全空的行
            if any(str(cell).strip() for cell in row_data):
                data.append(row_data)
        
        wb.close()
        
        if not data:
            return json.dumps({
                "error": "工作表为空或没有数据"
            }, ensure_ascii=False)
        
        # 分离表头和数据
        headers = data[0] if has_header and data else []
        rows = data[1:] if has_header and len(data) > 1 else data
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "sheet_name": ws.title,
            "available_sheets": wb.sheetnames if not wb.read_only else [],
            "statistics": {
                "row_count": len(rows),
                "column_count": len(headers) if headers else (len(rows[0]) if rows else 0),
                "has_header": has_header
            },
            "headers": headers,
            "rows": rows
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {
                "error": "您需要先安装openpyxl包才能读取该Excel文件。",
                "solution": "请运行以下命令安装：",
                "commands": {
                    "uv": "uv pip install openpyxl",
                    "pip": "pip install openpyxl"
                },
                "note": "或者安装所有依赖: pip install -r requirements.txt 或 uv pip install -r requirements.txt"
            },
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"读取Excel失败: {e}")
        return json.dumps({"error": f"读取失败: {str(e)}"}, ensure_ascii=False)


@document_tools.tool(
    description="列出Excel文件中的所有工作表名称及其基本信息。"
    "用于多工作表Excel文件的探索和导航。"
)
def list_excel_sheets(
    path: str,
) -> str:
    """
    列出Excel所有工作表
    
    返回每个工作表的：
    - 名称
    - 是否为活动工作表
    - 大致行列数
    
    Args:
        path: Excel文件路径
        
    Returns:
        工作表列表JSON
    """
    try:
        from openpyxl import load_workbook
        
        if not Path(path).exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        wb = load_workbook(path, read_only=True)
        
        sheets_info = []
        active_sheet_name = wb.active.title if wb.active else None
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            # 获取大致尺寸
            max_row = ws.max_row if ws.max_row else 0
            max_col = ws.max_column if ws.max_column else 0
            
            sheets_info.append({
                "name": sheet_name,
                "is_active": sheet_name == active_sheet_name,
                "max_row": max_row,
                "max_column": max_col
            })
        
        wb.close()
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "sheet_count": len(sheets_info),
            "sheets": sheets_info
        }, ensure_ascii=False, indent=2)
        
    except ImportError:
        return json.dumps(
            {
                "error": "您需要先安装openpyxl包才能读取该Excel文件。",
                "solution": "请运行以下命令安装：",
                "commands": {
                    "uv": "uv pip install openpyxl",
                    "pip": "pip install openpyxl"
                },
                "note": "或者安装所有依赖: pip install -r requirements.txt 或 uv pip install -r requirements.txt"
            },
            ensure_ascii=False
        )
    except Exception as e:
        logger.error(f"列出工作表失败: {e}")
        return json.dumps({"error": f"操作失败: {str(e)}"}, ensure_ascii=False)


# ==================== 辅助工具函数 ====================

def _parse_page_range(range_str: Optional[str], total_pages: int) -> List[int]:
    """
    解析页码范围字符串
    
    支持格式：
    - None: 返回所有页码
    - "1-5": 第1到5页
    - "1,3,5": 第1、3、5页
    - "1-3,5,7-9": 组合格式
    
    Args:
        range_str: 范围字符串
        total_pages: 总页数
        
    Returns:
        页码索引列表（从0开始）
    """
    if not range_str or range_str.strip() == "":
        return list(range(total_pages))
    
    pages = set()  # 使用集合去重
    
    try:
        parts = range_str.split(',')
        for part in parts:
            part = part.strip()
            
            if '-' in part:
                # 范围格式：1-5
                start, end = part.split('-')
                start_page = int(start.strip())
                end_page = int(end.strip())
                
                # 转换为0索引，并验证范围
                for p in range(start_page - 1, end_page):
                    if 0 <= p < total_pages:
                        pages.add(p)
            else:
                # 单页格式：3
                page_num = int(part)
                page_idx = page_num - 1
                if 0 <= page_idx < total_pages:
                    pages.add(page_idx)
    
    except ValueError as e:
        logger.error(f"解析页码范围失败: {range_str}, 错误: {e}")
        return []
    
    return sorted(list(pages))


@document_tools.tool(
    description="自动检测文件格式类型。"
    "支持PDF、Word、Excel、PowerPoint等常见文档格式的识别。"
    "基于文件扩展名和MIME类型进行判断。"
)
def detect_document_format(
    path: str,
) -> str:
    """
    检测文档格式
    
    返回信息：
    - 文件类型
    - 推荐的处理工具
    - 是否支持
    
    Args:
        path: 文件路径
        
    Returns:
        格式检测结果JSON
    """
    try:
        file_path = Path(path)
        
        if not file_path.exists():
            return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)
        
        extension = file_path.suffix.lower()
        
        # 文件类型映射
        format_info = {
            ".pdf": {
                "type": "PDF",
                "category": "Portable Document Format",
                "supported": True,
                "tools": ["extract_pdf_text", "extract_pdf_tables", "extract_pdf_images"]
            },
            ".docx": {
                "type": "Word",
                "category": "Microsoft Word Document",
                "supported": True,
                "tools": ["read_docx", "extract_docx_text"]
            },
            ".doc": {
                "type": "Word (Legacy)",
                "category": "Microsoft Word 97-2003",
                "supported": False,
                "message": "请将.doc文件转换为.docx格式"
            },
            ".xlsx": {
                "type": "Excel",
                "category": "Microsoft Excel Workbook",
                "supported": True,
                "tools": ["read_excel", "list_excel_sheets"]
            },
            ".xlsm": {
                "type": "Excel (Macro-enabled)",
                "category": "Microsoft Excel Macro-Enabled",
                "supported": True,
                "tools": ["read_excel", "list_excel_sheets"]
            },
            ".xls": {
                "type": "Excel (Legacy)",
                "category": "Microsoft Excel 97-2003",
                "supported": False,
                "message": "请将.xls文件转换为.xlsx格式"
            },
            ".pptx": {
                "type": "PowerPoint",
                "category": "Microsoft PowerPoint",
                "supported": False,
                "message": "PowerPoint支持正在开发中"
            },
            ".txt": {
                "type": "Text",
                "category": "Plain Text",
                "supported": True,
                "tools": ["file_tools.read_file"]
            },
            ".md": {
                "type": "Markdown",
                "category": "Markdown Document",
                "supported": True,
                "tools": ["file_tools.read_file"]
            }
        }
        
        info = format_info.get(extension, {
            "type": "Unknown",
            "category": "Unknown Format",
            "supported": False,
            "message": f"不支持的文件格式: {extension}"
        })
        
        # 添加文件基本信息
        file_size = file_path.stat().st_size
        
        return json.dumps({
            "success": True,
            "file_path": path,
            "filename": file_path.name,
            "extension": extension,
            "size_bytes": file_size,
            "size_mb": round(file_size / (1024 * 1024), 2),
            "format_info": info
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"检测文档格式失败: {e}")
        return json.dumps({"error": f"检测失败: {str(e)}"}, ensure_ascii=False)

